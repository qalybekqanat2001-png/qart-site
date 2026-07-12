import tempfile
from itertools import zip_longest

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from contract_content import BOLD_SPLIT, build_contract, build_requisites


def add_runs(paragraph, text: str, base_bold: bool = False, size: Pt | None = None):
    for part in BOLD_SPLIT.split(text):
        if not part:
            continue
        is_bold_span = part.startswith("**") and part.endswith("**")
        run = paragraph.add_run(part[2:-2] if is_bold_span else part)
        run.bold = True if is_bold_span else base_bold
        if size:
            run.font.size = size


def generate_contract_docx(data: dict) -> str:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    for kind, text in build_contract(data):
        para = doc.add_paragraph()

        if kind == "title":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_runs(para, text, base_bold=True, size=Pt(14))
        elif kind == "center":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_runs(para, text)
        elif kind == "heading":
            para.paragraph_format.space_before = Pt(10)
            para.paragraph_format.space_after = Pt(6)
            add_runs(para, text, base_bold=True, size=Pt(12))
        elif kind == "bullet":
            add_runs(para, text)
        elif kind == "empty":
            pass
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_runs(para, text)

    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(6)
    req = build_requisites(data)
    add_runs(heading, req["heading"], base_bold=True, size=Pt(12))

    rows = list(zip_longest(req["executor_rows"], req["client_rows"], fillvalue=""))
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = True
    for row_idx, (left, right) in enumerate(rows):
        left_cell = table.cell(row_idx, 0)
        right_cell = table.cell(row_idx, 1)
        left_cell.paragraphs[0].paragraph_format.space_after = Pt(2)
        right_cell.paragraphs[0].paragraph_format.space_after = Pt(2)
        add_runs(left_cell.paragraphs[0], left)
        add_runs(right_cell.paragraphs[0], right)

    doc.add_paragraph()
    sig_para = doc.add_paragraph()
    left_sig, right_sig = req["signature_row"]
    add_runs(sig_para, f"{left_sig}     {right_sig}")

    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    return tmp.name
